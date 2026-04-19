# -*- coding: utf-8 -*-
"""
word_engine.py — 基于 format_presets JSON 驱动的 Word 报告生成引擎
不再硬编码字体/字号/行距，统一从 JSON 预设中读取。
"""
import json, os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


_PRESET_DIR = os.path.join(os.path.dirname(__file__), '..', 'format_presets')


def load_preset(preset_name='thesis_songti'):
    """加载格式预设"""
    path = os.path.join(_PRESET_DIR, f'{preset_name}.json')
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


class ReportBuilder:
    """Word 报告构建器"""

    def __init__(self, preset_name='thesis_songti'):
        self.preset = load_preset(preset_name)
        self.doc = Document()
        self._setup_page()
        self._table_count = 0

    def _setup_page(self):
        """设置页面"""
        pg = self.preset['page']
        sec = self.doc.sections[0]
        sec.page_width = Cm(pg['width_cm'])
        sec.page_height = Cm(pg['height_cm'])
        sec.top_margin = Cm(pg['top_margin_cm'])
        sec.bottom_margin = Cm(pg['bottom_margin_cm'])
        sec.left_margin = Cm(pg['left_margin_cm'])
        sec.right_margin = Cm(pg['right_margin_cm'])

    def next_table_no(self):
        """自增表号"""
        self._table_count += 1
        return self._table_count

    @property
    def table_count(self):
        return self._table_count

    def add_title(self, text):
        """添加报告标题"""
        cfg = self.preset['heading']
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = cfg['font_en']
        rpr = run._element.get_or_add_rPr()
        rpr.get_or_add_rFonts().set(qn('w:eastAsia'), cfg['font_cn'])

    def add_heading(self, text, level=1):
        """添加标题"""
        cfg = self.preset['heading']
        h = self.doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = cfg['font_en']
            run.font.size = Pt(cfg['font_size_pt'])
            run.font.bold = cfg.get('bold', True)
            run.font.color.rgb = RGBColor.from_string(cfg['color'])
            rpr = run._element.get_or_add_rPr()
            rpr.get_or_add_rFonts().set(qn('w:eastAsia'), cfg['font_cn'])

    def _apply_line_spacing(self, pf, cfg):
        """应用行距设置：优先倍数行距，其次固定磅值"""
        if 'line_spacing_multiple' in cfg:
            # float 值 → python-docx 自动设为倍数行距（如 1.5 = 1.5倍）
            pf.line_spacing = cfg['line_spacing_multiple']
        elif 'line_spacing_pt' in cfg:
            pf.line_spacing = Pt(cfg['line_spacing_pt'])

    def _apply_first_indent(self, pf, cfg):
        """应用首行缩进：优先字符数（字符×字号），其次固定cm"""
        if 'first_line_indent_chars' in cfg:
            # 2字符 = 2 × 字号pt，确保不同字号下始终"2字符"
            pf.first_line_indent = Pt(cfg['font_size_pt'] * cfg['first_line_indent_chars'])
        elif 'first_line_indent_cm' in cfg:
            pf.first_line_indent = Cm(cfg['first_line_indent_cm'])

    def add_body_text(self, text, indent=True):
        """添加正文段落"""
        cfg = self.preset['body']
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(cfg.get('space_before_pt', 0))
        pf.space_after = Pt(cfg.get('space_after_pt', 0))
        pf.left_indent = Cm(cfg.get('left_indent_cm', 0))
        pf.right_indent = Cm(cfg.get('right_indent_cm', 0))
        self._apply_line_spacing(pf, cfg)
        if indent:
            self._apply_first_indent(pf, cfg)
        run = p.add_run(text)
        run.font.size = Pt(cfg['font_size_pt'])
        run.font.name = cfg['font_en']
        rpr = run._element.get_or_add_rPr()
        rpr.get_or_add_rFonts().set(qn('w:eastAsia'), cfg['font_cn'])

    def add_note(self, text):
        """添加注释"""
        cfg = self.preset['note']
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(cfg.get('space_before_pt', 0))
        pf.space_after = Pt(cfg.get('space_after_pt', 0))
        pf.left_indent = Cm(cfg.get('left_indent_cm', 0))
        pf.right_indent = Cm(cfg.get('right_indent_cm', 0))
        pf.first_line_indent = Cm(0)
        self._apply_line_spacing(pf, cfg)
        run = p.add_run(text)
        run.font.size = Pt(cfg['font_size_pt'])
        run.font.name = cfg['font_en']
        rpr = run._element.get_or_add_rPr()
        rpr.get_or_add_rFonts().set(qn('w:eastAsia'), cfg['font_cn'])

    def _set_cell_font(self, cell, text, bold=False,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER):
        """设置单元格字体 + 垂直居中"""
        cfg = self.preset['table']
        cell.text = ''
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = alignment
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(cfg.get('line_spacing_pt', 12))
        pf.first_line_indent = Cm(0)
        run = p.add_run(str(text))
        run.font.size = Pt(cfg['font_size_pt'])
        run.font.name = cfg['font_en']
        run.font.bold = bold
        rpr = run._element.get_or_add_rPr()
        rpr.get_or_add_rFonts().set(qn('w:eastAsia'), cfg['font_cn'])

    def add_three_line_table(self, headers, data_rows, title=''):
        """添加三线表"""
        tcfg = self.preset['table']
        ttcfg = self.preset['table_title']

        # 表标题
        if title:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(ttcfg['space_before_pt'])
            pf.space_after = Pt(ttcfg['space_after_pt'])
            pf.first_line_indent = Cm(0)
            run = p.add_run(title)
            run.font.size = Pt(ttcfg['font_size_pt'])
            run.font.name = ttcfg['font_en']
            run.font.bold = ttcfg['bold']
            rpr = run._element.get_or_add_rPr()
            rpr.get_or_add_rFonts().set(qn('w:eastAsia'), ttcfg['font_cn'])

        # 创建表格
        table = self.doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        for j, h in enumerate(headers):
            self._set_cell_font(table.rows[0].cells[j], h, bold=True)

        # 数据行
        for i, row in enumerate(data_rows):
            for j, val in enumerate(row):
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                self._set_cell_font(table.rows[i + 1].cells[j], val, alignment=align)

        # 三线表边框
        ns = nsdecls("w")
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {ns}/>')
        top_sz = tcfg.get('border_top_sz', 12)
        bot_sz = tcfg.get('border_bottom_sz', 12)
        hdr_sz = tcfg.get('border_header_bottom_sz', 4)

        tblPr.append(parse_xml(
            f'<w:tblBorders {ns}>'
            f'<w:top w:val="single" w:sz="{top_sz}" w:space="0" w:color="000000"/>'
            f'<w:bottom w:val="single" w:sz="{bot_sz}" w:space="0" w:color="000000"/>'
            f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
            f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
            f'</w:tblBorders>'))

        for cell in table.rows[0].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(parse_xml(
                f'<w:tcBorders {ns}>'
                f'<w:bottom w:val="single" w:sz="{hdr_sz}" w:space="0" w:color="000000"/>'
                f'</w:tcBorders>'))

        return table

    @staticmethod
    def set_table_col_widths(table, widths_cm):
        """设置表格列宽（厘米列表）

        用法: rb.set_table_col_widths(table, [4, 2, 2, 2, 2])
        """
        for row in table.rows:
            for i, width in enumerate(widths_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    def add_page_break(self):
        """插入分页符"""
        self.doc.add_page_break()

    def add_figure(self, image_path, caption='', width_cm=14):
        """插入图片 + 图题

        Args:
            image_path: 图片文件路径
            caption: 图题文字（如 '图1  中介效应模型'）
            width_cm: 图片宽度（厘米，默认14cm≈页面内容宽度）
        """
        # 图片
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(3)
        pf.first_line_indent = Cm(0)
        run = p.add_run()
        run.add_picture(image_path, width=Cm(width_cm))

        # 图题
        if caption:
            ttcfg = self.preset['table_title']  # 复用表标题格式
            p2 = self.doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf2 = p2.paragraph_format
            pf2.space_before = Pt(3)
            pf2.space_after = Pt(6)
            pf2.first_line_indent = Cm(0)
            run2 = p2.add_run(caption)
            run2.font.size = Pt(ttcfg['font_size_pt'])
            run2.font.name = ttcfg['font_en']
            run2.font.bold = ttcfg.get('bold', True)
            rpr = run2._element.get_or_add_rPr()
            rpr.get_or_add_rFonts().set(qn('w:eastAsia'), ttcfg['font_cn'])

    def save(self, path):
        """保存文档"""
        os.makedirs(os.path.dirname(path) or '.', exist_ok=True)
        self.doc.save(path)
        print(f'报告已保存: {path}')

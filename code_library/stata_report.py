# -*- coding: utf-8 -*-
"""
Stata outreg2 TXT 解析 + Word/Excel 报告生成模板
用途：读取 Stata outreg2 导出的 TXT 文件，生成三线表 Word 和 Excel
输入：stata_output/*.txt（outreg2格式 + file write 格式）
输出：交付成果/Stata回归结果表.docx, 交付成果/Stata分析结果.xlsx
日期：2026-03-25
参考项目：133 90/generate_stata_report.py

使用方式：
    1. 在 tables_order 中按顺序列出需要的表格
    2. 确保 stata_output/ 下有对应的 .txt 文件
    3. 运行本脚本即可生成 Word + Excel + zip
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

import os
import re
import shutil
import zipfile
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ══════ 配置 ══════
STATA_OUTPUT_DIR = 'stata_output'
OUT_DIR = '交付成果2'
WORD_PATH = os.path.join(OUT_DIR, 'Stata回归结果表.docx')
EXCEL_PATH = os.path.join(OUT_DIR, 'Stata分析结果.xlsx')
ZIP_PATH = '交付成果2.zip'

# 表格顺序：(TXT文件名不含扩展名, 表号前缀)
tables_order = [
    ('描述性统计', '表1 '),
    ('基准回归', '表2 '),
    ('中介效应', '表3 '),
    ('中介检验', '表3b '),
    ('内生性处理', '表4 '),
    ('稳健性检验', '表5 '),
    ('异质性分析', '表6 '),
]


# ══════ Word 格式工具函数 ══════
def set_cell_font(cell, text, font_cn='宋体', font_en='Times New Roman',
                  size=9, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """设置单元格字体（宋体+TNR 标准）"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = font_en
    rPr = run._element.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), font_cn)
    # 紧凑行距
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(
        f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" '
        f'w:line="280" w:lineRule="exact"/>'
    )
    pPr.append(spacing)


def set_three_line_border(table):
    """设置三线表边框（顶粗+底粗，无竖线）"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>'
    )
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    borders = parse_xml(borders_xml)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(borders)


def add_header_border(table, row_idx):
    """为表头行添加底部细线"""
    for cell in table.rows[row_idx].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tcBorders>'
        )
        for old in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old)
        tcPr.append(borders)


# ══════ outreg2 TXT 解析 ══════
def parse_outreg2_txt(filepath):
    """解析 outreg2 生成的 TXT 文件，返回 (title, headers, data_rows)"""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = [l.rstrip('\r\n') for l in f.readlines()]

    if not lines:
        return '', [], []

    title = lines[0].strip() if lines else ''
    headers = []
    data_rows = []

    for i, line in enumerate(lines):
        cols = line.split('\t')
        if i == 0:
            continue  # 标题行
        elif i <= 2:
            # 表头行（列名）
            if any(c.strip() for c in cols[1:]):
                headers = [c.strip() for c in cols]
        else:
            if any(c.strip() for c in cols):
                data_rows.append([c.strip() for c in cols])

    return title, headers, data_rows


# ══════ 添加表格到 Word ══════
def add_table_to_doc(doc, title, headers, data_rows, table_prefix=''):
    """将解析后的数据添加为三线表到 Word 文档"""
    # 标题
    t_para = doc.add_paragraph()
    t_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    display_title = table_prefix + title if title else table_prefix + '回归结果'
    t_run = t_para.add_run(display_title)
    t_run.font.size = Pt(10)
    t_run.font.bold = True
    t_run.font.name = 'Times New Roman'
    rPr = t_run._element.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '黑体')

    if not data_rows:
        return

    n_cols = max(len(headers), max(len(r) for r in data_rows))
    n_rows = len(data_rows) + 1  # +1 for header
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 填充表头
    for j in range(min(len(headers), n_cols)):
        align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
        set_cell_font(table.rows[0].cells[j], headers[j], bold=True, alignment=align)

    # 填充数据
    for i, row_data in enumerate(data_rows):
        for j in range(min(len(row_data), n_cols)):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(table.rows[i + 1].cells[j], row_data[j], alignment=align)

    set_three_line_border(table)
    add_header_border(table, 0)

    # 注释
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = note.add_run('注：括号内为标准误；*** p<0.01, ** p<0.05, * p<0.1。')
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()


# ══════ 主流程 ══════
def main():
    os.makedirs(OUT_DIR, exist_ok=True)

    doc = Document()
    excel_data = {}

    for name, prefix in tables_order:
        txt_path = os.path.join(STATA_OUTPUT_DIR, f'{name}.txt')
        if not os.path.exists(txt_path):
            print(f'⚠ 跳过 {name}（文件不存在）')
            continue

        title, headers, data_rows = parse_outreg2_txt(txt_path)
        add_table_to_doc(doc, title, headers, data_rows, prefix)

        # Excel
        if headers and data_rows:
            max_cols = max(len(headers), max(len(r) for r in data_rows))
            headers_padded = headers + [''] * (max_cols - len(headers))
            rows_padded = [r + [''] * (max_cols - len(r)) for r in data_rows]
            excel_data[name] = pd.DataFrame(rows_padded, columns=headers_padded)

        print(f'✓ {prefix}{name}')

    doc.save(WORD_PATH)
    print(f'\nWord 已保存: {WORD_PATH}')

    if excel_data:
        with pd.ExcelWriter(EXCEL_PATH, engine='openpyxl') as w:
            for sn, df in excel_data.items():
                df.to_excel(w, sheet_name=sn[:31], index=False)
        print(f'Excel 已保存: {EXCEL_PATH}')

    # 打包 zip
    if os.path.exists(ZIP_PATH):
        os.remove(ZIP_PATH)
    with zipfile.ZipFile(ZIP_PATH, 'w', zipfile.ZIP_DEFLATED) as z:
        for root, dirs, files in os.walk(OUT_DIR):
            for file in files:
                full = os.path.join(root, file)
                z.write(full, os.path.relpath(full, '.'))
    print(f'\n打包完成: {ZIP_PATH}')
    print('全部完成!')


if __name__ == '__main__':
    main()

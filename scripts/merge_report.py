# -*- coding: utf-8 -*-
"""
用途：合并多个独立分析文档为一份完整报告
输入文件：交付成果/01_*.docx, 02_*.docx, ...
输出文件：交付成果/分析报告（合并版）.docx
日期：2026-03-30（修复图片丢失 + 分节符问题）
"""
import sys
import os
import re
import glob

sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'code_library'))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from copy import deepcopy
from lxml import etree


# ══════ 图片关系迁移工具 ══════

def _copy_image_parts(src_doc, dst_doc, element):
    """
    将 element 中引用的所有图片（blip）从 src_doc 迁移到 dst_doc。
    处理流程：
      1. 在 element XML 中查找所有 a:blip 节点的 r:embed 属性
      2. 通过源文档的 part.rels 找到对应的 image part
      3. 将 image part 添加到目标文档的 part 中，获取新的 rId
      4. 更新 element 中的 r:embed 为新 rId
    """
    nsmap = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    }
    blips = element.findall('.//a:blip[@r:embed]', nsmap)
    if not blips:
        return

    src_part = src_doc.part
    dst_part = dst_doc.part

    for blip in blips:
        old_rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if not old_rid:
            continue
        try:
            rel = src_part.rels[old_rid]
            image_part = rel.target_part
        except (KeyError, AttributeError):
            continue

        # 添加图片 part 到目标文档并获取新 rId
        new_rid = dst_part.relate_to(image_part, RT.IMAGE)
        blip.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', new_rid)


def _copy_ole_parts(src_doc, dst_doc, element):
    """复制嵌入的 OLE 对象（如 Excel 图表等）"""
    nsmap = {
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    }
    # 查找 oleObject 和其他引用了 r:id 的嵌入元素
    for node in element.iter():
        rid_attr = node.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        if not rid_attr:
            continue
        try:
            rel = src_doc.part.rels[rid_attr]
            target_part = rel.target_part
            new_rid = dst_doc.part.relate_to(target_part, rel.reltype)
            node.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', new_rid)
        except (KeyError, AttributeError):
            pass


# ══════ 核心合并逻辑 ══════

def merge_docx_files(input_dir, output_path, title='数据分析报告'):
    """
    合并多个独立分析 .docx 文件为一份完整报告

    规则：
    1. 按文件名前缀数字排序（01_xxx.docx, 02_xxx.docx, ...）
    2. 自动重新编号表格（表1, 表2, ... 全局连续）
    3. 自动重新编号图片（图1, 图2, ...）
    4. 保留三线表格式、字体、段落样式
    5. 正确迁移图片二进制数据（修复图片丢失 bug）
    6. 文件间使用分页符分隔，不插入多余分节符
    """
    # 1. 收集所有 .docx 文件，按名称排序
    pattern = os.path.join(input_dir, '*.docx')
    files = sorted(glob.glob(pattern))

    # 排除合并版自身
    files = [f for f in files if '合并版' not in os.path.basename(f)
             and '合并' not in os.path.basename(f)]

    if not files:
        print(f'❌ 未找到 .docx 文件: {input_dir}')
        return

    print(f'═══ 合并文档 ═══\n')
    print(f'输入目录: {input_dir}')
    print(f'找到 {len(files)} 个文件:')
    for f in files:
        print(f'  {os.path.basename(f)}')

    # 2. 创建主文档
    merged_doc = Document()

    # 设置默认样式
    style = merged_doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5

    # 添加总标题
    title_para = merged_doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = '黑体'
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. 逐文件合并
    table_counter = 0
    figure_counter = 0
    WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    for file_idx, filepath in enumerate(files):
        filename = os.path.basename(filepath)
        print(f'\n处理: {filename}')

        try:
            sub_doc = Document(filepath)
        except Exception as e:
            print(f'  ⚠️ 无法打开: {e}')
            continue

        img_count = 0
        is_first_element = True  # 标记每个子文档的首个段落

        # 复制段落和表格
        for element in sub_doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

            if tag == 'sectPr':
                # 跳过节属性——合并文档使用统一的页面设置
                continue

            # 深拷贝 XML 节点
            new_element = deepcopy(element)

            # 在每个子文档的第一个段落上注入 pageBreakBefore（第一个文件除外）
            if is_first_element and file_idx > 0 and tag == 'p':
                pPr = new_element.find(f'{{{WNS}}}pPr')
                if pPr is None:
                    pPr = etree.SubElement(new_element, f'{{{WNS}}}pPr')
                    new_element.insert(0, pPr)
                page_break = etree.SubElement(pPr, f'{{{WNS}}}pageBreakBefore')
                is_first_element = False

            # 迁移图片和嵌入对象的二进制数据
            _copy_image_parts(sub_doc, merged_doc, new_element)
            _copy_ole_parts(sub_doc, merged_doc, new_element)

            if tag == 'tbl':
                table_counter += 1

            # 统计图片数量
            nsmap_a = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
            if new_element.findall('.//a:blip', nsmap_a):
                img_count += len(new_element.findall('.//a:blip', nsmap_a))

            # 插入到 body 的最后一个 sectPr 之前（确保位置正确）
            body_sectPr = merged_doc.element.body.find(f'{{{WNS}}}sectPr')
            if body_sectPr is not None:
                body_sectPr.addprevious(new_element)
            else:
                merged_doc.element.body.append(new_element)

        print(f'  ✅ 段落/表格已合并，包含 {img_count} 张图片')

    # 4. 全局重新编号表格（两阶段：先建映射，再批量替换）
    old_to_new_table = {}  # 旧编号 -> 新编号
    table_num = 0
    for para in merged_doc.paragraphs:
        text = para.text.strip()
        # 只匹配表标题段落（以"表X"开头的段落）
        m = re.match(r'^表\s*(\d+)', text)
        if m:
            old_num = int(m.group(1))
            table_num += 1
            old_to_new_table[old_num] = table_num

    # 批量替换所有段落中的表编号（标题+正文引用）
    if old_to_new_table:
        def _replace_table_num(match):
            old = int(match.group(1))
            new = old_to_new_table.get(old, old)
            return f'表{new}'

        for para in merged_doc.paragraphs:
            if re.search(r'表\s*\d+', para.text):
                for run in para.runs:
                    if re.search(r'表\s*\d+', run.text):
                        run.text = re.sub(r'表\s*(\d+)', _replace_table_num, run.text)

    # 5. 全局重新编号图片（两阶段：先建映射，再批量替换）
    old_to_new_fig = {}
    fig_num = 0
    for para in merged_doc.paragraphs:
        text = para.text.strip()
        m = re.match(r'^图\s*(\d+)', text)
        if m:
            old_num = int(m.group(1))
            fig_num += 1
            old_to_new_fig[old_num] = fig_num

    if old_to_new_fig:
        def _replace_fig_num(match):
            old = int(match.group(1))
            new = old_to_new_fig.get(old, old)
            return f'图{new}'

        for para in merged_doc.paragraphs:
            if re.search(r'图\s*\d+', para.text):
                for run in para.runs:
                    if re.search(r'图\s*\d+', run.text):
                        run.text = re.sub(r'图\s*(\d+)', _replace_fig_num, run.text)

    # 6. 保存合并文档
    merged_doc.save(output_path)
    print(f'\n✅ 合并完成: {output_path}')
    print(f'   共 {table_counter} 个表格，{fig_num} 张图片')
    print(f'   来自 {len(files)} 个文件')


def list_step_files(input_dir):
    """列出所有分步文档及其状态"""
    pattern = os.path.join(input_dir, '*.docx')
    files = sorted(glob.glob(pattern))
    files = [f for f in files if '合并' not in os.path.basename(f)]

    print(f'═══ 分步文档状态 ═══\n')
    for f in files:
        size = os.path.getsize(f)
        print(f'  ✅ {os.path.basename(f)} ({size//1024}KB)')

    if not files:
        print('  ℹ️ 暂无分步文档')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法:')
        print('  合并: python merge_report.py "交付成果/" ["报告标题"]')
        print('  列表: python merge_report.py "交付成果/" --list')
        sys.exit(1)

    input_dir = sys.argv[1]

    if '--list' in sys.argv:
        list_step_files(input_dir)
    else:
        title = sys.argv[2] if len(sys.argv) > 2 else '数据分析报告'
        output_path = os.path.join(input_dir, f'分析报告（合并版）.docx')
        merge_docx_files(input_dir, output_path, title)
